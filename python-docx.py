# !/usr/bin/env python3
# _*_coding:utf-8_*_
# __author__：FLS





from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.shared import Inches
from docx.oxml.ns import qn

document = Document()


document.add_heading()
run=document.add_heading('中国人民银行关于《标准化票据管理办法（征求意见稿）》公开征求意见的通知',level=1).add_run(u'应用场景')
run.font.name=u'微软雅黑'
run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')






document.save('demo.docx')




# 一级标题：小标宋 居中 二号 行间距32磅
# 二级标题：居中 黑体 三号 行间距32磅 段前/后 0.3行
# 三级标题：楷体加粗 三号 行间距29.3磅
# 四级标题：楷体 三号 行间距29.3磅
# 正文：仿宋 三号 行间距29.3磅



