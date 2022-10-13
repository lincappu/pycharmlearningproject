'''
只支持ms的docx文件来操作，不支持doc的文件，
官网： https://python-docx.readthedocs.io/en/latest/index.html
'''

import  docx

document=docx.Document('a.docx')


document.add_heading('ROLE',level=1)
paragraph=document.add_paragraph('Test file')
document.add_paragraph('这是一个样式为list的段落',style='ListBullet')



document.save('a.docx')